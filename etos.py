#!/usr/bin/env python3
from deep_translator import GoogleTranslator
import docx2txt
import docx
import sys


def over_five_thousand(word_list):
    """Feed text into google translate in blocks under the 5000 character limit"""
    j = 0
    words = ""
    t_list = []
    for i, word in enumerate(word_list):
        j += len(word) + 2  # Add total charcter count to j variable

        if j < 5000:  # If character count is under 5000 add word to words
            words += word + " "

            # On last word in the list if j < 5000
            # translate remaining words and add to list
            if i == (len(word_list) - 1):
                t_list.append(
                    GoogleTranslator(source="en", target="es").translate(words)
                )

        else:
            # Translate 5000 character block of text and add to list
            t_list.append(GoogleTranslator(source="en", target="es").translate(words))
            j = 0  # Reset character counter
            words = word + " "

            # Translate and add last word to list if j > 5000 on the last word
            if i == (len(word_list) - 1):
                t_list.append(
                    GoogleTranslator(source="en", target="es").translate(word)
                )

    return t_list


def docx_translate(name):
    """Translate a docx file to spanish and save it to another docx file"""
    try:
        text = docx2txt.process(name)
    except FileNotFoundError:
        print("File Not Found")
        exit(-1)

    length = len(text)
    if length == 0:
        print("No text in file")
        exit(-1)

    # Remove empty lines from docx file, append lines to line list with newline
    line = []
    for i in text.splitlines():
        if i != "":
            line.append(i + "\n")
    lines = "".join(line)

    if length < 5000:
        sp_text = GoogleTranslator(source="en", target="es").translate(lines)
        spanish_doc = docx.Document()
        spanish_doc.add_paragraph(sp_text)
        spanish_doc.save("spanish " + name)

    else:
        word_list = lines.split(" ")  # Split lines into list of words
        t_list = over_five_thousand(word_list)
        spanish_doc = docx.Document()

        spanish_doc.add_paragraph("\n".join(t_list))
        spanish_doc.save("spanish " + name)


def txt_translate(name):
    try:
        file = open(name, "r")
    except FileNotFoundError:
        print("File not found")
        exit(-1)

    text = file.read()
    length = len(text)

    if length == 0:
        print("No text in file")
        exit(-1)
    line = file.readlines()
    word_list = text.split(" ")

    file.close()
    if length < 5000:
        sp_text = GoogleTranslator(source="en", target="es").translate("\n".join(line))
        with open("spanish " + name, "w") as file:
            file.write(sp_text)
    else:
        t_list = over_five_thousand(word_list)
        with open("spanish " + name, "w") as file:
            file.write("\n".join(t_list))


try:  # Check command for a valid file
    name = sys.argv[1]
except IndexError:
    print("Error!: Must enter a file")
    print("Example: etos file or etos /path/to/file")
    exit()

if name == "help" or name == "--help":  # Return syntax if user types help
    print("Example: etos file or etos /path/to/file")
    exit()

if name[-5:] == ".docx":
    docx_translate(name)

elif name[-4:] == ".txt":
    txt_translate(name)

else:
    print("Unable to translate this filetype")
